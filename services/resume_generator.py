import io
import re
from html import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (CondPageBreak, KeepTogether, Paragraph,
                                 SimpleDocTemplate, Spacer)

# Matches role lines: "Title | 01/2024 – Present" or "Title [01/2024 – 12/2025]"
_ROLE_RE = re.compile(
    r'(\|\s*\d{2}/\d{4})'           # | MM/YYYY
    r'|(\[\s*\d{2}/\d{4})'          # [MM/YYYY
    r'|(\b\d{4}\s*[–—-]\s*(\d{4}|Present|Current)\b)',  # YYYY – YYYY/Present
    re.IGNORECASE,
)


def _is_role_line(s: str) -> bool:
    if s.startswith("•") or s.startswith("-"):
        return False
    return bool(_ROLE_RE.search(s))


def _label_lines(text: str) -> list[tuple[str, str]]:
    """Return (original_line, label) pairs.

    Labels: name | contact | section | company | role | bullet | empty | body
    """
    raw = text.split("\n")
    labels = ["body"] * len(raw)

    # Index non-empty lines in order
    non_empty = [(i, raw[i].strip()) for i in range(len(raw)) if raw[i].strip()]

    for rank, (i, s) in enumerate(non_empty):
        if rank == 0:
            labels[i] = "name"
        elif rank == 1 and not (s.isupper() and len(s) > 2):
            labels[i] = "contact"
        elif s.isupper() and len(s) > 2:
            labels[i] = "section"
        elif _is_role_line(s):
            labels[i] = "role"
        elif s.startswith("•") or (s.startswith("-") and len(s) > 2):
            labels[i] = "bullet"

    # Mark empty lines
    for i, line in enumerate(raw):
        if not line.strip():
            labels[i] = "empty"

    # Second pass: body lines immediately before a role line → company
    non_empty_labels = [(i, labels[i]) for i, _ in non_empty]
    for rank, (i, lbl) in enumerate(non_empty_labels):
        if lbl != "body":
            continue
        if rank + 1 < len(non_empty_labels):
            _, next_lbl = non_empty_labels[rank + 1]
            if next_lbl == "role":
                labels[i] = "company"

    return list(zip(raw, labels))


# ── Public API ────────────────────────────────────────────────────────────────

def generate_docx(resume_text: str, original_bytes: bytes | None = None,
                  original_filename: str = "") -> bytes:
    if original_bytes and original_filename.lower().endswith(".docx"):
        try:
            return _docx_from_template(resume_text, original_bytes)
        except Exception:
            pass
    return _docx_basic(resume_text)


def generate_pdf(resume_text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=inch, rightMargin=inch,
    )
    doc.build(_build_pdf_story(resume_text))
    return buf.getvalue()


# ── DOCX: template-based ──────────────────────────────────────────────────────

def _docx_from_template(resume_text: str, template_bytes: bytes) -> bytes:
    template = Document(io.BytesIO(template_bytes))
    styles = _extract_docx_styles(template)

    doc = Document(io.BytesIO(template_bytes))
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    _populate_docx(doc, resume_text, styles)

    if sect_pr is not None and body.find(qn("w:sectPr")) is None:
        body.append(sect_pr)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_docx_styles(doc: Document) -> dict:
    base_font = "Calibri"
    base_size = 10.5

    # Try to detect base font from body paragraphs
    for para in doc.paragraphs:
        s = para.text.strip()
        if not s or s.isupper():
            continue
        for run in para.runs:
            if run.font.name:
                base_font = run.font.name
            if run.font.size:
                base_size = run.font.size.pt
            break
        break

    margins = {"top": Inches(0.75), "bottom": Inches(0.75),
               "left": Inches(1.0), "right": Inches(1.0)}
    if doc.sections:
        s = doc.sections[0]
        margins = {"top": s.top_margin, "bottom": s.bottom_margin,
                   "left": s.left_margin, "right": s.right_margin}

    # Name size: scan first paragraph
    name_size = 18.0
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if non_empty:
        r = next((r for r in non_empty[0].runs if r.text.strip()), None)
        if r and r.font.size:
            name_size = r.font.size.pt

    return {
        "font":     base_font,
        "name_size":    name_size,
        "contact_size": max(8.0, base_size - 1),
        "section_size": base_size + 0.5,
        "body_size":    base_size,
        "margins":  margins,
    }


def _populate_docx(doc: Document, resume_text: str, styles: dict):
    font      = styles["font"]
    margins   = styles["margins"]

    for section in doc.sections:
        section.top_margin    = margins["top"]
        section.bottom_margin = margins["bottom"]
        section.left_margin   = margins["left"]
        section.right_margin  = margins["right"]

    for line, label in _label_lines(resume_text):
        s = line.strip()

        if label == "empty":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            continue

        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(2)

        if label == "name":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_after = Pt(0)
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["name_size"])
            run.font.bold = True

        elif label == "contact":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_after = Pt(6)
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["contact_size"])

        elif label == "section":
            fmt.space_before = Pt(10)
            fmt.space_after  = Pt(4)
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["section_size"])
            run.font.bold = True

        elif label == "company":
            fmt.space_before = Pt(6)
            fmt.space_after  = Pt(0)
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["body_size"])
            run.font.bold = True
            run.font.underline = True

        elif label == "role":
            fmt.space_after = Pt(2)
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["body_size"])
            run.font.bold = True

        elif label == "bullet":
            fmt.left_indent  = Inches(0.25)
            fmt.space_after  = Pt(1)
            text = s.lstrip("•- ").strip()
            run = p.add_run(f"• {text}")
            run.font.name = font
            run.font.size = Pt(styles["body_size"])

        else:  # body
            run = p.add_run(s)
            run.font.name = font
            run.font.size = Pt(styles["body_size"])


def _docx_basic(resume_text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    default_styles = {
        "font": "Calibri", "name_size": 18.0, "contact_size": 10.0,
        "section_size": 11.0, "body_size": 10.5,
        "margins": {"top": Inches(0.75), "bottom": Inches(0.75),
                    "left": Inches(1.0), "right": Inches(1.0)},
    }
    _populate_docx(doc, resume_text, default_styles)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _build_pdf_story(resume_text: str) -> list:
    base = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "RName", parent=base["Normal"],
        fontSize=18, fontName="Helvetica-Bold",
        alignment=1, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "RContact", parent=base["Normal"],
        fontSize=10, fontName="Helvetica",
        alignment=1, spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "RSection", parent=base["Normal"],
        fontSize=11, fontName="Helvetica-Bold",
        spaceBefore=10, spaceAfter=4,
    )
    company_style = ParagraphStyle(
        "RCompany", parent=base["Normal"],
        fontSize=10.5, fontName="Helvetica-Bold",
        spaceBefore=6, spaceAfter=0,
    )
    role_style = ParagraphStyle(
        "RRole", parent=base["Normal"],
        fontSize=10, fontName="Helvetica-Bold",
        spaceBefore=0, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "RBody", parent=base["Normal"],
        fontSize=10, fontName="Helvetica",
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "RBullet", parent=base["Normal"],
        fontSize=10, fontName="Helvetica",
        leftIndent=16, spaceAfter=1,
    )

    labeled = _label_lines(resume_text)

    # Split into page-break groups: header block + one block per section
    groups: list[list] = []
    current_group: list = []

    for line, label in labeled:
        s = line.strip()

        if label == "empty":
            current_group.append(Spacer(1, 3))
            continue

        if label == "name":
            current_group.append(Paragraph(escape(s), name_style))
        elif label == "contact":
            current_group.append(Paragraph(escape(s), contact_style))
        elif label == "section":
            # Start a new group at each section header
            if current_group:
                groups.append(current_group)
            current_group = [Paragraph(escape(s), section_style)]
        elif label == "company":
            current_group.append(Paragraph(f"<u>{escape(s)}</u>", company_style))
        elif label == "role":
            current_group.append(Paragraph(escape(s), role_style))
        elif label == "bullet":
            text = escape(s.lstrip("•- ").strip())
            current_group.append(Paragraph(f"• {text}", bullet_style))
        else:
            current_group.append(Paragraph(escape(s), body_style))

    if current_group:
        groups.append(current_group)

    story: list = []
    for i, group in enumerate(groups):
        if i > 0:
            story.append(CondPageBreak(1.5 * inch))
        story.append(KeepTogether(group))

    return story
